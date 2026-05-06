"""
Word-eksport af analyser og svarbreve.

Konverterer Claudes markdown-output til pænt formaterede .docx-filer med
overskrifter, fed tekst, kursiv og punkttegn. Returnerer bytes som
Streamlit kan tilbyde som download.

Claude returnerer typisk tekst med:
  - # Overskrift eller ## Underoverskrift
  - **fed tekst**
  - *kursiv tekst*
  - - punkttegn
  - Normale afsnit adskilt af tomme linjer
"""

import re
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Regex til at finde markdown-formatering i en linje
INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)([^*]+?)\*(?!\*)")


def _parse_inline(tekst):
    """
    Splitter en tekstlinje op i segmenter med formatering:
    [("fed", "tekst"), ("normal", " og "), ("kursiv", "kursiv")]
    """
    # Vi finder bold først (gør dem til markører), så italic på resten.
    segmenter = []
    rest = tekst

    # Scan gennem teksten og find matches
    pos = 0
    while pos < len(rest):
        bold_match = INLINE_BOLD.search(rest, pos)
        italic_match = INLINE_ITALIC.search(rest, pos)

        # Find det tidligste match
        kandidater = [m for m in [bold_match, italic_match] if m is not None]
        if not kandidater:
            segmenter.append(("normal", rest[pos:]))
            break

        naeste = min(kandidater, key=lambda m: m.start())
        if naeste.start() > pos:
            segmenter.append(("normal", rest[pos:naeste.start()]))

        if naeste is bold_match:
            segmenter.append(("fed", naeste.group(1)))
        else:
            segmenter.append(("kursiv", naeste.group(1)))
        pos = naeste.end()

    return segmenter


def _tilfoej_formateret_linje(paragraph, tekst):
    """Tilføjer tekst til et Word-afsnit med bold/italic fra markdown."""
    for stil, segment in _parse_inline(tekst):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        if stil == "fed":
            run.bold = True
        elif stil == "kursiv":
            run.italic = True


def _er_overskrift(linje):
    """Returnerer (niveau, ren_tekst) hvis linjen er en overskrift, ellers None."""
    m = re.match(r"^(#{1,3})\s+(.+)", linje)
    if m:
        return len(m.group(1)), m.group(2).strip()
    return None


def _er_fed_overskrift(linje):
    """
    Genkender '**1. Indledning**' og lignende som overskrift.
    Claude bruger ofte dette format frem for # markdown-overskrifter.
    """
    m = re.match(r"^\*\*(.+?)\*\*\s*$", linje.strip())
    if m:
        indhold = m.group(1).strip()
        # Hvis det ligner en overskrift (starter med tal eller er kort), brug det
        if re.match(r"^\d+\.", indhold) or len(indhold) < 70:
            return indhold
    return None


def _er_bullet(linje):
    """Returnerer (niveau, ren_tekst) hvis det er et punkttegn, ellers None."""
    m = re.match(r"^(\s*)[-•*]\s+(.+)", linje)
    if m:
        niveau = len(m.group(1)) // 2
        return niveau, m.group(2).strip()
    return None


def markdown_til_docx_bytes(markdown_tekst, titel="Dokument", undertitel=None):
    """
    Konverterer Claudes markdown-output til en pænt formateret .docx-fil.
    Returnerer bytes.
    """
    doc = Document()

    # Dokumentopsætning — standardfont og margener
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Sektionsopsætning — rimelige margener
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titel
    titel_p = doc.add_heading(titel, level=0)
    titel_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Undertitel/metadata
    if undertitel:
        meta_p = doc.add_paragraph(undertitel)
        meta_p.runs[0].italic = True
        meta_p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Dato
    dato_p = doc.add_paragraph(
        f"Udarbejdet: {datetime.now().strftime('%d-%m-%Y kl. %H:%M')}"
    )
    dato_p.runs[0].italic = True
    dato_p.runs[0].font.size = Pt(9)
    dato_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # blank linje

    # Parse markdown linje for linje
    linjer = markdown_tekst.split("\n")
    i = 0
    while i < len(linjer):
        linje = linjer[i].rstrip()

        # Tom linje → blank paragraph
        if not linje.strip():
            # Undgå dobbelte blanke afsnit
            i += 1
            continue

        # Almindelig markdown overskrift (# eller ##)
        overskrift = _er_overskrift(linje)
        if overskrift:
            niveau, tekst = overskrift
            doc.add_heading(tekst, level=min(niveau, 3))
            i += 1
            continue

        # '**1. Indledning**' som overskrift
        fed_overskrift = _er_fed_overskrift(linje)
        if fed_overskrift:
            doc.add_heading(fed_overskrift, level=2)
            i += 1
            continue

        # Punkttegn
        bullet = _er_bullet(linje)
        if bullet:
            _, tekst = bullet
            p = doc.add_paragraph(style="List Bullet")
            _tilfoej_formateret_linje(p, tekst)
            i += 1
            continue

        # Ellers: normalt afsnit. Saml alle sammenhængende ikke-tomme linjer
        # (der ikke er overskrift/bullet) som ét afsnit.
        afsnit_linjer = [linje]
        j = i + 1
        while j < len(linjer):
            n = linjer[j].rstrip()
            if not n.strip():
                break
            if _er_overskrift(n) or _er_fed_overskrift(n) or _er_bullet(n):
                break
            afsnit_linjer.append(n)
            j += 1
        # Join med mellemrum (hårde linjeskift i markdown er normalt soft wraps)
        tekst = " ".join(afsnit_linjer)
        p = doc.add_paragraph()
        _tilfoej_formateret_linje(p, tekst)
        i = j

    # Konvertér til bytes
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def analyse_til_docx(spoergsmaal, svar, klage_filnavn=None):
    """Konverterer en AI-analyse til en .docx-fil."""
    undertitel_dele = [f"Spørgsmål: {spoergsmaal}"]
    if klage_filnavn:
        undertitel_dele.insert(0, f"Klage: {klage_filnavn}")
    return markdown_til_docx_bytes(
        svar,
        titel="Juridisk analyse",
        undertitel=" — ".join(undertitel_dele),
    )


def _saet_paragraf_under_streg(paragraph):
    """
    Tegner en tynd vandret streg UNDER det angivne afsnit. Bruges på
    'Vedr.'-linjen for at give den den klassiske brevhoved-streg.
    """
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")        # 1pt streg (8/8)
    bottom.set(qn("w:space"), "4")     # afstand til tekst
    bottom.set(qn("w:color"), "000000")
    p_borders.append(bottom)
    p_pr.append(p_borders)


def _fjern_celle_kant(cell):
    """Sætter alle 4 kanter på en tabelcelle til 'nil' så cellen er
    usynlig. Bruges på header-tabellen så modtager + logo ser ud som
    fritlæggende tekst (ikke som en synlig tabel)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for kant in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{kant}")
        b.set(qn("w:val"), "nil")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _byg_bilag_liste(doc, bilag_liste):
    """
    Renders 'Bilag:'-overskrift efterfulgt af en pæn 2-kolonne-tabel
    (uden synlige kanter) der viser hvert bilag som "Bilag X | beskrivelse".
    Layoutet matcher det Mikkel sendte i referencebilledet:

        Bilag:

        Bilag A    TUIs bemærkninger til sagen
        Bilag B    Første bekræftelsesmail på rejsebestillingen
        Bilag C    ...

    bilag_liste er en liste af dicts på formen:
      [{"bogstav": "A", "overskrift": "TUIs bemærkninger til sagen"}, ...]

    Hvis listen er tom rendres ingenting (defensivt).
    """
    if not bilag_liste:
        return

    # "Bilag:"-overskrift
    bilag_p = doc.add_paragraph()
    bilag_run = bilag_p.add_run("Bilag:")
    bilag_run.bold = True
    bilag_run.font.size = Pt(12)

    # 2-kolonne-tabel uden synlige kanter
    tabel = doc.add_table(rows=len(bilag_liste), cols=2)
    tabel.autofit = False
    for r, post in enumerate(bilag_liste):
        bogstav_celle = tabel.cell(r, 0)
        beskrivelse_celle = tabel.cell(r, 1)
        bogstav_celle.width = Cm(2.5)
        beskrivelse_celle.width = Cm(13.5)
        _fjern_celle_kant(bogstav_celle)
        _fjern_celle_kant(beskrivelse_celle)

        # Venstre: "Bilag X"
        bogstav_p = bogstav_celle.paragraphs[0]
        bogstav_run = bogstav_p.add_run(f"Bilag {post.get('bogstav', '')}")
        bogstav_run.font.size = Pt(11)

        # Højre: beskrivelsen
        beskrivelse_p = beskrivelse_celle.paragraphs[0]
        beskrivelse_run = beskrivelse_p.add_run(
            post.get("overskrift") or ""
        )
        beskrivelse_run.font.size = Pt(11)

    # Lidt luft før brødteksten
    doc.add_paragraph()


def _byg_svarbrev_header(
    doc, profil, sagsnummer, klagers_navn, hoeringssvar_nr,
    bilag_liste=None,
):
    """
    Bygger den klassiske brevhoved-opsætning til et svarbrev:

      [PAKKEREJSE-ANKENÆVNET     ]   [logo top højre]
      [Haldor Topsøes Alle 1, ...]
      [2800 Kgs. Lyngby          ]

                             By, DD-MM-YYYY

      Vedr.: Sag nr. XXXX – Klagers navn, N. høringssvar
      ───────────────────────────────────────────────────

    Felter der mangler (fx hvis sagsnummer ikke kunne udledes) springes
    over uden at crashe. Logoet springes over hvis filen ikke findes.
    """
    from selskab_profiler import hent_logo_sti, hent_by

    # ---------- 2-kolonne header: modtager-adresse + logo ----------
    header_tabel = doc.add_table(rows=1, cols=2)
    header_tabel.autofit = False
    # Bredde: ca. 60% adresse, 40% logo (totalt ~16cm = A4 minus margener)
    venstre_celle = header_tabel.cell(0, 0)
    hoejre_celle = header_tabel.cell(0, 1)
    venstre_celle.width = Cm(10)
    hoejre_celle.width = Cm(6)
    _fjern_celle_kant(venstre_celle)
    _fjern_celle_kant(hoejre_celle)

    # Venstre: Pakkerejse-Ankenævnets adresse (samme for alle selskaber)
    adresse_p = venstre_celle.paragraphs[0]
    fed_run = adresse_p.add_run("PAKKEREJSE-ANKENÆVNET")
    fed_run.bold = True
    fed_run.font.size = Pt(11)
    venstre_celle.add_paragraph("Haldor Topsøes Alle 1, Bygning 91")
    venstre_celle.add_paragraph("2800 Kgs. Lyngby")

    # Højre: logo (hvis fil findes — ellers tom celle)
    logo_p = hoejre_celle.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_sti = hent_logo_sti(profil)
    if logo_sti:
        try:
            logo_p.add_run().add_picture(logo_sti, width=Cm(3.5))
        except Exception as e:
            # Defensivt: hvis billedet er korrupt eller i ukendt format,
            # spring det over så svarbrevet stadig kan downloades.
            print(f"DEBUG: kunne ikke indsætte logo {logo_sti}: {e}")

    # ---------- By + dato (højrejusteret, ca. 2 blanke linjer nede) ----------
    doc.add_paragraph()  # luft
    by = hent_by(profil) or ""
    dato_str = datetime.now().strftime("%d-%m-%Y")
    by_dato_text = (
        f"{by}, {dato_str}" if by else dato_str
    )
    by_dato_p = doc.add_paragraph()
    by_dato_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    by_dato_run = by_dato_p.add_run(by_dato_text)
    by_dato_run.font.size = Pt(11)

    # ---------- "Vedr."-linje med vandret streg under ----------
    doc.add_paragraph()  # luft før Vedr-linjen
    vedr_dele = ["Vedr.: "]
    if sagsnummer:
        vedr_dele.append(f"Sag nr. {sagsnummer}")
    if klagers_navn:
        if sagsnummer:
            vedr_dele.append(f" – {klagers_navn}")
        else:
            vedr_dele.append(klagers_navn)
    if hoeringssvar_nr:
        vedr_dele.append(f", {hoeringssvar_nr}. høringssvar")

    vedr_p = doc.add_paragraph()
    vedr_run = vedr_p.add_run("".join(vedr_dele))
    vedr_run.bold = True
    vedr_run.font.size = Pt(11)
    _saet_paragraf_under_streg(vedr_p)

    # Lidt ekstra luft før brødteksten
    doc.add_paragraph()

    # NOTE: Bilag-listen rendres IKKE længere her. Den er flyttet til
    # bunden af svarbrevet (efter "Med venlig hilsen" + signatur).
    # Se svarbrev_til_docx for placeringen.


def svarbrev_til_docx(
    svarbrev,
    klage_filnavn=None,
    sagsnummer="",
    klagers_navn="",
    hoeringssvar_nr=1,
    profil=None,
    bilag_liste=None,
):
    """
    Konverterer et svarbrev til en .docx-fil med klassisk brev-opsætning:
    modtager-adresse øverst til venstre, selskabs-logo øverst til højre,
    by + dato højrejusteret, derefter en 'Vedr.'-linje med vandret streg.

    Selve brødteksten (svarbrevet) er den AI-genererede markdown og
    formateres med samme inline-parser som de øvrige docx-eksporter.

    Parametre:
      svarbrev          — den AI-genererede markdown
      klage_filnavn     — bruges KUN til at navngive download-filen
                          udadtil; står ikke i selve brevet
      sagsnummer        — fx "25-109-8024327" (kan være tom)
      klagers_navn      — fx "Laura Stephanie Uhler" (kan være tom)
      hoeringssvar_nr   — 1, 2 eller 3
      profil            — selskabs-profil-dict (fra selskab_profiler).
                          Hvis None bruges den aktive profil.
      bilag_liste       — liste af dicts på formen
                          [{"bogstav": "A", "overskrift": "..."}, ...]
                          Vises som bilag-oversigt nederst i svarbrevet
                          (efter Samlet vurdering, før venlig hilsen).
                          Bilag A er pr. konvention selve svarbrevet.
                          Hvis None/tom rendres ingen bilag-liste.
    """
    # Lazy-import så modulet stadig kan importeres i miljøer hvor
    # selskab_profiler endnu ikke er på plads (defensivt — bør altid være der)
    try:
        from selskab_profiler import hent_aktiv_profil
        if profil is None:
            profil = hent_aktiv_profil()
    except Exception as e:
        print(f"DEBUG: kunne ikke hente selskabs-profil: {e}")
        profil = {"navn": "", "by": "", "logo_fil": None}

    doc = Document()

    # Standard-font og margener (samme som markdown_til_docx_bytes så
    # brødteksten ser ens ud)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---------- BREVHOVED ----------
    _byg_svarbrev_header(
        doc=doc,
        profil=profil,
        sagsnummer=sagsnummer or "",
        klagers_navn=klagers_navn or "",
        hoeringssvar_nr=hoeringssvar_nr,
        bilag_liste=bilag_liste or [],
    )

    # ---------- BRØDTEKST (svarbrevet) ----------
    # Genbrug samme markdown-parser som markdown_til_docx_bytes, men
    # tilføj indholdet til vores eksisterende doc i stedet for at lave
    # et nyt dokument. Vi inliner parser-loopet her for at undgå at
    # opfinde et nyt API på markdown_til_docx_bytes.
    linjer = (svarbrev or "").split("\n")
    i = 0
    while i < len(linjer):
        linje = linjer[i].rstrip()

        if not linje.strip():
            i += 1
            continue

        overskrift = _er_overskrift(linje)
        if overskrift:
            niveau, tekst = overskrift
            doc.add_heading(tekst, level=min(niveau, 3))
            i += 1
            continue

        fed_overskrift = _er_fed_overskrift(linje)
        if fed_overskrift:
            doc.add_heading(fed_overskrift, level=2)
            i += 1
            continue

        bullet = _er_bullet(linje)
        if bullet:
            _, tekst = bullet
            p = doc.add_paragraph(style="List Bullet")
            _tilfoej_formateret_linje(p, tekst)
            i += 1
            continue

        # Almindeligt afsnit: saml sammenhængende linjer
        afsnit_linjer = [linje]
        j = i + 1
        while j < len(linjer):
            n = linjer[j].rstrip()
            if not n.strip():
                break
            if (
                _er_overskrift(n)
                or _er_fed_overskrift(n)
                or _er_bullet(n)
            ):
                break
            afsnit_linjer.append(n)
            j += 1
        tekst = " ".join(afsnit_linjer)
        p = doc.add_paragraph()
        _tilfoej_formateret_linje(p, tekst)
        i = j

    # ---------- BILAG-LISTE I BUNDEN ----------
    # Bilag-listen rendres som det allersidste i brevet — efter
    # "Med venlig hilsen" og signaturen som AI'en har skrevet i selve
    # brødteksten. Ekstra spacer-paragraph for visuel separation.
    if bilag_liste:
        doc.add_paragraph()  # spacer
        _byg_bilag_liste(doc, bilag_liste)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------- PDF-EKSPORT (bruges primært til anonymiserede bilag) ----------

def markdown_til_pdf_bytes(markdown_tekst, titel="Dokument", undertitel=None):
    """
    Konverterer en markdown-tekst til en simpel, læsbar PDF.
    Returnerer bytes. Bruges fx til anonymiserede bilag hvor brugeren
    foretrækker PDF over Word.

    Reportlab importes lokalt så evt. miljøer uden reportlab stadig kan
    importere modulet (eksport til docx bliver ikke påvirket).
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,
    )

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.4 * cm,
        bottomMargin=2.4 * cm,
        title=titel,
    )

    styles = getSampleStyleSheet()
    titel_style = ParagraphStyle(
        "DocTitel",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        spaceAfter=6,
    )
    undertitel_style = ParagraphStyle(
        "DocUndertitel",
        parent=styles["Italic"],
        fontSize=10,
        textColor="#64748B",
        spaceAfter=18,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        spaceAfter=8,
        alignment=TA_LEFT,
    )
    h2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        spaceBefore=12,
        spaceAfter=6,
    )
    h3_style = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        spaceBefore=10,
        spaceAfter=4,
    )

    flow = []
    flow.append(Paragraph(_escape_html(titel), titel_style))
    if undertitel:
        flow.append(Paragraph(_escape_html(undertitel), undertitel_style))

    linjer = (markdown_tekst or "").split("\n")
    bullet_buffer = []

    def flush_bullets():
        if not bullet_buffer:
            return
        items = [
            ListItem(Paragraph(_md_inline_til_html(b), body_style))
            for b in bullet_buffer
        ]
        flow.append(ListFlowable(
            items,
            bulletType="bullet",
            start="•",
            leftIndent=18,
            bulletFontName="Helvetica",
            bulletFontSize=10,
        ))
        bullet_buffer.clear()

    for raa in linjer:
        linje = raa.rstrip()
        if not linje.strip():
            flush_bullets()
            flow.append(Spacer(1, 4))
            continue

        if linje.startswith("# "):
            flush_bullets()
            flow.append(Paragraph(_md_inline_til_html(linje[2:]), titel_style))
        elif linje.startswith("## "):
            flush_bullets()
            flow.append(Paragraph(_md_inline_til_html(linje[3:]), h2_style))
        elif linje.startswith("### "):
            flush_bullets()
            flow.append(Paragraph(_md_inline_til_html(linje[4:]), h3_style))
        elif linje.lstrip().startswith(("- ", "* ")):
            content = linje.lstrip()[2:]
            bullet_buffer.append(content)
        else:
            flush_bullets()
            flow.append(Paragraph(_md_inline_til_html(linje), body_style))

    flush_bullets()

    doc.build(flow)
    return buf.getvalue()


def _escape_html(s):
    """Minimal HTML-escape til reportlab Paragraph."""
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _md_inline_til_html(tekst):
    """
    Konverterer markdown inline-formatering til de mini-HTML-tags som
    reportlab's Paragraph forstår. Understøtter **fed** og *kursiv*.
    """
    tekst = _escape_html(tekst)
    tekst = INLINE_BOLD.sub(r"<b>\1</b>", tekst)
    tekst = INLINE_ITALIC.sub(r"<i>\1</i>", tekst)
    return tekst

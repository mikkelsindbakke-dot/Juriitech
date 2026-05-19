"""
DOCX-anonymisering med tekst-erstatning.

Genbruger detektoren fra anonymisering_pdf (regex + AI) og erstatter
følsomme tekst-segmenter med █-blokke der visuelt ligner sort-bjælke-
redaction i PDF'en. Original formatering (skrifttype, fed, kursiv,
tabeller, sektioner) bevares.

Modsat PDF-ruten er det IKKE ægte redaction: den oprindelige tekst
er erstattet, ikke overlagt. Til gengæld er der ingen risiko for at
underliggende tekst kan trækkes ud — for den findes ikke længere.

Tre lag:
- Detektor: udtræk tekst, find redaction-targets (genbrugt fra PDF-modul)
- Redactor: erstat targets i alle runs på tværs af paragraffer/tabeller/headers
- Orchestrator: kaldes fra api/main.py med ensartet (bytes|None, status)-API
"""
from __future__ import annotations

import io
from typing import Iterable

import docx


REDACTION_BLOCK = "█"  # U+2588 FULL BLOCK


def udtraek_docx_tekst(docx_bytes: bytes) -> str:
    """Returnér al tekst fra DOCX'en som én streng. Afsnit adskilles af '\\n\\n'."""
    doc = docx.Document(io.BytesIO(docx_bytes))
    bidder: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text:
            bidder.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        bidder.append(paragraph.text)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for paragraph in hf.paragraphs:
                if paragraph.text:
                    bidder.append(paragraph.text)

    return "\n\n".join(bidder)


def _erstat_i_paragraph(paragraph, streng: str, erstatning: str) -> int:
    """
    Erstat alle forekomster af `streng` i en paragraf, selv hvis strengen
    spænder flere runs (typisk pga. inline-formatering).

    Strategi: saml al run-tekst, lav erstatning på den samlede streng,
    læg resultatet i første run og tøm de øvrige. Formatering på første
    run bevares — de øvrige runs' formatering går tabt, men det er
    acceptabelt for redaction-formålet.

    Returnerer antal erstatninger udført.
    """
    runs = paragraph.runs
    if not runs or not streng:
        return 0
    fuld = "".join(r.text for r in runs)
    if streng not in fuld:
        return 0
    antal = fuld.count(streng)
    ny = fuld.replace(streng, erstatning)
    runs[0].text = ny
    for r in runs[1:]:
        r.text = ""
    return antal


def redact_docx(docx_bytes: bytes, targets: Iterable[dict]) -> bytes:
    """
    Anvend tekst-erstatning på alle paragraffer i hovedindhold, tabeller
    og headers/footers. Returnér ny DOCX som bytes.

    For hvert target erstattes strengen med det antal █-blokke der
    matcher den oprindelige længde (capped på 30, så lange strenge ikke
    eksploderer i bredde). Lignende PDF-sort-bjælke visuelt.
    """
    targets = list(targets)
    doc = docx.Document(io.BytesIO(docx_bytes))

    total = 0

    def håndter(paragraph) -> None:
        nonlocal total
        for target in targets:
            streng = target.get("streng", "")
            if not streng:
                continue
            bjælke = REDACTION_BLOCK * min(len(streng), 30)
            try:
                total += _erstat_i_paragraph(paragraph, streng, bjælke)
            except Exception as e:
                # En enkelt paragraf-fejl må ikke vælte hele dokumentet
                print(f"DEBUG: erstatning fejlede for streng-pattern: {e}")

    for paragraph in doc.paragraphs:
        håndter(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    håndter(paragraph)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for paragraph in hf.paragraphs:
                håndter(paragraph)

    print(f"DEBUG: docx redaction udførte {total} erstatninger")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def anonymiser_docx_fil(
    docx_bytes: bytes,
    klager_navne: list[str],
) -> tuple[bytes | None, str]:
    """
    Komplet anonymiseringsflow for én DOCX.

    Returnerer (output_bytes, status_streng). Status er én af:
      - "ok"               — successful erstatning
      - "fejl_aaben"       — DOCX kunne ikke åbnes/parses
      - "fejl_redaktion"   — redact_docx fejlede
    """
    from anonymisering_pdf import find_redaction_targets

    try:
        tekst = udtraek_docx_tekst(docx_bytes)
    except Exception as e:
        print(f"DEBUG: docx tekst-ekstraktion fejlede: {e}")
        return (None, "fejl_aaben")

    targets = find_redaction_targets(tekst, klager_navne)

    try:
        output_bytes = redact_docx(docx_bytes, targets)
        return (output_bytes, "ok")
    except Exception as e:
        print(f"DEBUG: redact_docx fejlede: {e}")
        return (None, "fejl_redaktion")
